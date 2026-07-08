"""ProposalForge Agent — Copilot-Studio-style UI.

One combined Overview page (Details + Knowledge on the left, Test/Chat on the
right) plus Agents, Evaluation, Analytics and Test Suites. No login. Answers are
grounded in the uploaded knowledge, guardrailed for PII/PHI, and the chat
suggests follow-ups based on the previous question. Reads PDFs/images including
architecture diagrams & charts via Azure Document Intelligence + LLM vision.
"""
from __future__ import annotations

import os
import re
import io
import csv
import time
import uuid
from datetime import datetime
from pathlib import Path

import streamlit as st

from src.agents.llm_backend import get_llm_client
from src.retrieval.retriever import get_retriever
from src.common.file_processor import FileProcessor
from src.common import doc_analysis, guardrails as gr
from src.agents.compliance_matrix import build_matrix
from src.common.load_tester import run_load_test
from src.common.tracking_store import TrackingStore
from src.common.test_runner import run_test_suite
from src.retrieval.ragas_eval import evaluate_ragas
from src.retrieval.ragas_llm import evaluate_metrics, METRIC_KEYS

# --------------------------------------------------------------------------- #
st.set_page_config(page_title="ProposalForge Agent", page_icon="🤖",
                   layout="wide", initial_sidebar_state="collapsed")

# ------------------------------- colourful theme --------------------------- #
st.markdown("""
<style>
:root{--ink:#16233B;--teal:#0F9AA8;--tealD:#0B7A86;--amber:#E9950C;--violet:#5B5BD6;}
.block-container{padding-top:1.2rem;}
/* centered gradient title banner */
.pf-hero{background:linear-gradient(120deg,#16233B 0%,#0F9AA8 60%,#3FBFD4 100%);
  color:#fff;border-radius:18px;padding:22px 26px;text-align:center;margin-bottom:14px;
  box-shadow:0 8px 24px rgba(15,154,168,.25);}
.pf-hero h1{margin:0;font-size:34px;font-weight:800;letter-spacing:.3px;}
.pf-hero p{margin:.35rem 0 0;font-size:15px;opacity:.95;}
.pf-badge{display:inline-block;padding:3px 12px;border-radius:999px;font-weight:700;font-size:13px;}
.pf-badge.on{background:#DCFCE7;color:#166534;} .pf-badge.off{background:#FEF3C7;color:#92400E;}
.pf-dom{display:inline-block;padding:2px 10px;border-radius:999px;background:#E8F6F8;color:#0B7A86;font-weight:700;font-size:12px;}
/* coloured tab bar */
.stTabs [data-baseweb="tab-list"]{gap:8px;}
.stTabs [data-baseweb="tab"]{background:#EEF3F8;border-radius:12px 12px 0 0;padding:10px 18px;font-weight:700;color:#33415C;}
.stTabs [aria-selected="true"]{background:linear-gradient(135deg,#0F9AA8,#3FBFD4);color:#fff;}
/* pill buttons */
div[data-testid="stButton"]>button{border-radius:999px;border:1px solid #CFE0E3;font-weight:600;}
div[data-testid="stButton"]>button:hover{border-color:#0F9AA8;color:#0F9AA8;}
h3,h4{color:#16233B;}
</style>
""", unsafe_allow_html=True)

# ------------------------------- resources --------------------------------- #
@st.cache_resource
def get_client(provider: str = "claude"):
    return get_llm_client(provider)          # (client, backend) with failover

@st.cache_resource
def get_guard():
    return gr.Guardrails()


@st.cache_resource
def get_tracker():
    """Persistent audit/tracking store — Azure SQL in prod, SQLite locally."""
    return TrackingStore()

@st.cache_resource
def get_engine():
    from src.orchestrator.engine import Engine
    return Engine()

DOMAINS = {
    "Finance": ["loan", "credit", "invest", "revenue", "bank", "interest", "capital",
                "portfolio", "risk", "subsidy", "budget", "tax", "roi", "cash"],
    "Healthcare": ["patient", "clinical", "diagnos", "treatment", "hospital", "phi",
                   "medical", "disease", "drug", "therapy", "care", "health"],
    "Agriculture": ["crop", "farm", "soil", "irrigation", "yield", "pest", "harvest",
                    "seed", "fertiliz", "livestock", "agri", "millet"],
}

def classify_domain(text: str) -> str:
    t = (text or "").lower()
    best, score = "General", 0
    for d, kws in DOMAINS.items():
        s = sum(1 for k in kws if k in t)
        if s > score:
            best, score = d, s
    return best

# ------------------------------- session ----------------------------------- #
def _init():
    ss = st.session_state
    ss.setdefault("documents", [])
    ss.setdefault("history", [])            # [{role,content,confidence,sources,domain,blocked}]
    ss.setdefault("retriever", None)
    ss.setdefault("retriever_backend", "none")
    ss.setdefault("agent_name", "ProposalForge Agent")
    ss.setdefault("agent_description",
                  "Answers questions about your proposals & documents across "
                  "Finance, Healthcare and Agriculture — grounded, cited and safe.")
    ss.setdefault("agent_instructions", "")
    ss.setdefault("guardrail_hits", 0)
    ss.setdefault("audit", [])
    ss.setdefault("uid", "user")
    ss.setdefault("sid", uuid.uuid4().hex[:8])
_init()

provider = st.session_state.get("llm_provider", "Claude")
client, backend = get_client(provider.lower())
guard = get_guard()
online = bool(getattr(client, "online", False))

# ------------------------------- helpers ----------------------------------- #
def rebuild_index():
    docs = [{"id": d["id"], "name": d["name"], "text": d["text"]}
            for d in st.session_state.documents]
    try:
        r = get_retriever()
        r.build_documents(docs)
        st.session_state.retriever = r
        st.session_state.retriever_backend = getattr(r, "backend", "none")
    except Exception as e:
        st.session_state.retriever = None
        st.session_state.retriever_backend = "none"
        st.warning(f"Vector index unavailable ({e}); using plain-text context.")

def audit(action: str, detail: str = ""):
    entry = {"time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
             "user": st.session_state.uid, "action": action, "detail": str(detail)[:180]}
    st.session_state.audit.insert(0, entry)
    st.session_state.audit = st.session_state.audit[:300]
    # Persist to the tracking DB (Azure SQL in prod, SQLite locally).
    try:
        get_tracker().log(action, st.session_state.uid, st.session_state.sid,
                          str(detail)[:300])
    except Exception:
        pass

def _audit_csv(rows) -> str:
    buf = io.StringIO()
    w = csv.DictWriter(buf, fieldnames=["time", "user", "action", "detail"])
    w.writeheader(); w.writerows(rows)
    return buf.getvalue()

_LABELS = {"OBJECTIVE": "Objective", "CURRENT SOLUTIONS": "Current solutions",
           "PROPOSED SOLUTIONS": "Proposed solutions", "INSIGHTS": "Insights"}

def _parse_sections(raw: str) -> dict:
    out = {v: "" for v in _LABELS.values()}
    cur = None
    for line in raw.splitlines():
        s = line.strip()
        head = s.upper().rstrip(":").strip()
        if head in _LABELS:
            cur = _LABELS[head]; continue
        if cur and s:
            out[cur] += (("\n" if out[cur] else "") + s.lstrip("-•* ").strip())
    return out

def analyze_fields(text: str) -> dict:
    """Objective / Current solutions / Proposed solutions / Insights for a doc."""
    if online and hasattr(client, "complete"):
        try:
            raw = client.complete(
                "You analyze business and proposal documents. Be concise.",
                "Extract these four sections as short bullet lines using these EXACT "
                "headers:\nOBJECTIVE:\nCURRENT SOLUTIONS:\nPROPOSED SOLUTIONS:\nINSIGHTS:\n\n"
                f"Document:\n{text[:6000]}")
            if raw and "OBJECTIVE" in raw.upper():
                return _parse_sections(raw)
        except Exception:
            pass
    # offline / fallback via the structured analyzer
    a = {}
    try:
        a = client.analyze_document(text[:6000])
    except Exception:
        a = {}
    fmt = lambda v: "\n".join(v) if isinstance(v, list) else (v or "")
    return {"Objective": fmt(a.get("objective", "")),
            "Current solutions": fmt(a.get("challenges", [])),
            "Proposed solutions": fmt(a.get("proposed_solutions", [])),
            "Insights": fmt(a.get("insights", []))}

def ingest_files(files, scan_visuals: bool):
    added = 0
    names = {d["name"] for d in st.session_state.documents}
    for uf in files:
        if uf.name in names:
            continue
        try:
            raw = bytes(uf.getbuffer())
            ext = os.path.splitext(uf.name)[1].lower()
            is_img = ext in (".png", ".jpg", ".jpeg")
            # text: local for docs; Azure DI / vision for images or when scanning
            if is_img:
                text = ""
            else:
                p = Path("temp"); p.mkdir(exist_ok=True)
                fp = p / uf.name; fp.write_bytes(raw)
                text = FileProcessor.extract_text(str(fp))
            visuals = {"available": False, "descriptions": []}
            if scan_visuals or is_img:
                da = doc_analysis.analyze_document(raw, uf.name, include_visuals=True)
                visuals = da.get("visuals", visuals)
                if not text:
                    text = da.get("text") or "\n\n".join(
                        f"[Figure p{d['page']}] {d['description']}"
                        for d in visuals.get("descriptions", []))
            dom = classify_domain(text or uf.name)
            st.session_state.documents.append({
                "id": uuid.uuid4().hex, "name": uf.name, "text": text or "",
                "domain": dom, "visuals": visuals,
                "analysis": analyze_fields(text or uf.name),
                "chars": len(text or ""), "at": datetime.now().strftime("%H:%M")})
            audit("KNOWLEDGE_UPLOAD",
                  f"{uf.name} · topic={dom} · {len(text or '')} chars"
                  + (" · visuals scanned" if (scan_visuals or is_img) else ""))
            added += 1
        except Exception as e:
            st.error(f"❌ {uf.name}: {e}")
    if added:
        rebuild_index()
    return added

ENGINE_DOMAIN_MAP = {"Finance": "finance", "Healthcare": "healthcare",
                     "Agriculture": "agriculture"}


def _ensure_general_domain(engine):
    """Register a permissive 'general' domain so multi-domain uploads have a home."""
    if "general" not in engine.domains:
        from src.common.config import DomainConfig
        engine.domains["general"] = DomainConfig(
            key="general", label="General", emoji="🗂️",
            persona=("You are a proposal solutions consultant. Answer grounded in the "
                     "retrieved documents across any business domain and give concrete "
                     "proposed solutions and insights."),
            routing_keywords=[], guardrails={}, clarify={}, follow_ups=[],
            cross_domain_hint=[])


def feed_engine_docs():
    """Feed uploaded docs (as text) into the LangGraph Engine store under mapped domains."""
    added = 0
    try:
        eng = get_engine(); _ensure_general_domain(eng)
        fed = st.session_state.setdefault("_engine_fed", set())
        for d in st.session_state.documents:
            if d["id"] in fed:
                continue
            text = (d.get("text") or "").strip()
            if not text:
                fed.add(d["id"]); continue
            dom = ENGINE_DOMAIN_MAP.get(d.get("domain", "General"), "general")
            name = (os.path.splitext(d["name"])[0] or d["name"]) + ".txt"
            try:
                added += eng.add_document(text.encode("utf-8"), name, dom)
                fed.add(d["id"])
            except Exception:
                pass
        st.session_state["_engine_fed"] = fed
    except Exception:
        pass
    return added


def followups(last_q: str, last_a: str):
    try:
        raw = client.get_auto_suggestions((last_a or "")[:1500], last_q or "")
        out = [s for s in raw if not guard.check_pii(s).triggered]
        return out[:3]
    except Exception:
        return []

def _grounded(answer: str, context: str) -> float:
    a = set(re.findall(r"[a-z]{4,}", (answer or "").lower()))
    c = set(re.findall(r"[a-z]{4,}", (context or "").lower()))
    return round(len(a & c) / max(1, len(a)), 2) if a else 0.0


SYSTEM_SOLUTIONS = (
    "You are ProposalForge, a senior proposal solutions consultant. The user has "
    "uploaded documents that describe business problems, current situations and "
    "proposals. Answer the user's question directly and specifically. When they ask "
    "about the problems or how to solve them, give concrete PROPOSED SOLUTIONS and "
    "actionable INSIGHTS to solve those problems, grounded in the uploaded documents. "
    "Be specific and practical; prefer short paragraphs and bullet points. If the "
    "documents don't cover something, say so briefly and give your best expert advice."
)

_REPORT_WORDS = ("download", "report", "generate a doc", "generate a document",
                 "proposed solutions doc", "solutions document", "word doc",
                 "docx", "give me a doc", "export", "downloadable")


def _report_intent(q: str) -> bool:
    ql = (q or "").lower()
    return any(w in ql for w in _REPORT_WORDS)


def _llm_solutions_for(doc: dict) -> dict:
    """Objective/Problems/Proposed solutions/Insights for one doc (LLM if online)."""
    an = doc.get("analysis", {}) or {}
    base = {"Objective": an.get("Objective", ""),
            "Problems / current situation": an.get("Current solutions", ""),
            "Proposed solutions & improvements": an.get("Proposed solutions", ""),
            "Insights": an.get("Insights", "")}
    if online and hasattr(client, "complete"):
        try:
            raw = client.complete(
                SYSTEM_SOLUTIONS,
                "For the document below, write improvement recommendations using these "
                "EXACT headers as short bullet lines:\nOBJECTIVE:\nPROBLEMS:\n"
                "PROPOSED SOLUTIONS:\nINSIGHTS:\n\nDocument:\n" + (doc.get("text", "")[:6000]))
            if raw and "PROPOSED SOLUTIONS" in raw.upper():
                # reuse the section parser (maps to Objective/Current/Proposed/Insights)
                p = _parse_sections(
                    raw.replace("PROBLEMS:", "CURRENT SOLUTIONS:"))
                return {"Objective": p.get("Objective", ""),
                        "Problems / current situation": p.get("Current solutions", ""),
                        "Proposed solutions & improvements": p.get("Proposed solutions", ""),
                        "Insights": p.get("Insights", "")}
        except Exception:
            pass
    return base


def build_solutions_report():
    """Build an in-memory .docx of proposed solutions for every uploaded doc."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    docs = st.session_state.documents
    d = Document()
    h = d.add_heading("Proposed Solutions & Improvement Recommendations", level=0)
    d.add_paragraph(f"Prepared by {st.session_state.agent_name} · "
                    f"{datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True
    d.add_paragraph(
        "This report reviews the uploaded proposal documents, summarises the problems "
        "they raise, and recommends concrete proposed solutions and insights to address "
        "them.")
    for i, doc in enumerate(docs, 1):
        d.add_heading(f"{i}. {doc['name']}  ({doc.get('domain','General')})", level=1)
        sec = _llm_solutions_for(doc)
        for label in ("Objective", "Problems / current situation",
                      "Proposed solutions & improvements", "Insights"):
            val = (sec.get(label) or "").strip()
            d.add_heading(label, level=2)
            if val:
                for ln in val.splitlines():
                    ln = ln.strip().lstrip("-•* ").strip()
                    if ln:
                        d.add_paragraph(ln, style="List Bullet")
            else:
                d.add_paragraph("—")
    buf = io.BytesIO()
    d.save(buf)
    return "Proposed_Solutions_Report.docx", buf.getvalue()

def answer_query(q: str):
    ss = st.session_state
    gin = guard.check_pii(q)
    ss.history.append({"role": "user", "content": q})
    audit("QUERY", q)
    if gin.triggered:
        ss.guardrail_hits += 1
        audit("GUARDRAIL_BLOCK", f"input · {gin.message}")
        ss.history.append({"role": "assistant", "blocked": True,
            "content": f"🛡️ **Guardrail — PII/PHI protection.** {gin.message}. "
                       "I can discuss the documents in general terms but won't reveal "
                       "personal or health identifiers."})
        return

    # ---- report-download intent -> build a proposed-solutions .docx ----
    if _report_intent(q) and ss.documents:
        try:
            fname, data = build_solutions_report()
            ss["report_file"] = (fname, data)
            audit("REPORT_DOWNLOAD", f"{fname} · {len(ss.documents)} docs")
            ss.history.append({"role": "assistant",
                "content": ("✅ I've prepared a **Proposed Solutions & Improvement "
                            "Recommendations** report covering your uploaded "
                            f"{len(ss.documents)} document(s) — objectives, the problems "
                            "raised, and concrete proposed solutions and insights for "
                            "each. Use the **Download** button just below the chat."),
                "confidence": 0.95, "sources": [d["name"] for d in ss.documents],
                "domain": "General", "grounded": 1.0})
        except Exception as e:
            ss.history.append({"role": "assistant",
                "content": f"I couldn't build the report ({e}). Please try again."})
        return

    # ---- retrieve grounded context ----
    sources, retrieved = [], []
    r = ss.retriever
    if r is not None:
        try:
            retrieved = r.search(q, top_k=4)
        except Exception:
            retrieved = []
    if retrieved:
        context = "\n\n".join(f"[Source: {h['source']}]\n{h['text']}" for h in retrieved)
        for h in retrieved:
            if h["source"] not in sources:
                sources.append(h["source"])
    elif ss.documents:
        context = ss.documents[-1]["text"][:3500]
        sources = [ss.documents[-1]["name"]]
    else:
        context = ""

    recent = [m for m in ss.history if m["role"] in ("user", "assistant")
              and not m.get("blocked")][-4:]
    convo = ""
    if recent:
        convo = "\n".join(f"{m['role'].capitalize()}: {m['content'][:280]}"
                          for m in recent)
    instr = (ss.agent_instructions or "").strip()

    # ---- answer: solution-oriented, DIRECT ----
    ans, conf = "", 0.0
    prompt = ""
    if instr:
        prompt += f"Agent instructions (follow these):\n{instr}\n\n"
    if convo:
        prompt += f"Conversation so far:\n{convo}\n\n"
    prompt += (f"Documents:\n{context}\n\n---\nUser question: {q}\n\n"
               "Answer directly. If it's about problems or solutions, give concrete "
               "proposed solutions and insights grounded in the documents.")
    if online and hasattr(client, "complete"):
        try:
            ans = client.complete(SYSTEM_SOLUTIONS, prompt) or ""
        except Exception:
            ans = ""
    if not ans:
        # offline / fallback — synthesise a direct answer from the analysis
        resp = client.answer_question(context, q,
                                      allow_general_knowledge=not bool(ss.documents))
        ans = resp.get("answer", "") or ""
        conf = float(resp.get("confidence", 0.0))
        wants_sol = any(w in q.lower() for w in
                        ("solution", "solve", "improve", "recommend", "fix", "how",
                         "problem", "address", "what should"))
        if ss.documents and wants_sol:
            an = ss.documents[-1].get("analysis", {})
            probs = (an.get("Current solutions") or "").strip()
            ps = (an.get("Proposed solutions") or "").strip()
            ins = (an.get("Insights") or "").strip()
            bull = lambda t: "\n".join(f"- {l.lstrip('-•* ').strip()}"
                                       for l in t.splitlines() if l.strip())
            if ps or ins:
                parts = ["Based on your uploaded documents, here are the problems "
                         "and how to address them:"]
                if probs:
                    parts.append("**Problems / current situation**\n" + bull(probs))
                if ps:
                    parts.append("**Proposed solutions**\n" + bull(ps))
                if ins:
                    parts.append("**Insights**\n" + bull(ins))
                ans = "\n\n".join(parts)     # lead with the solutions, not a not-found note
    conf = conf or (0.85 if (sources and ans) else 0.6 if ans else 0.0)

    gout = guard.check_pii(ans)
    if gout.triggered:
        ss.guardrail_hits += 1
        audit("GUARDRAIL_REDACT", "answer contained PII/PHI — redacted")
        ans = gr.redact_pii(ans)
    ss.history.append({"role": "assistant", "content": ans or
                       "I don't have enough in the documents to answer that yet.",
                       "confidence": conf, "sources": sources,
                       "domain": classify_domain(q + " " + ans),
                       "grounded": _grounded(ans, context)})

# ------------------------------- centered title ---------------------------- #
st.markdown(
    '<div class="pf-hero"><h1>🤖 ProposalForge Agent</h1>'
    '<p>Multi-document agent for any domain — grounded, cited & guardrailed · '
    'reads text, tables & diagrams</p></div>', unsafe_allow_html=True)

tab_over, tab_agents, tab_compliance, tab_eval, tab_analytics, tab_activity, tab_tests = st.tabs(
    ["🏠 Overview", "🤖 Agents", "📋 Compliance", "🧪 Evaluation",
     "📊 Analytics", "🧾 Activity", "✅ Test Suites"])

# =========================== OVERVIEW ====================================== #
with tab_over:
    st.markdown(f"### {st.session_state.agent_name}  "
                + (f'<span class="pf-badge on">● Ready · {backend}</span>' if online
                   else '<span class="pf-badge off">● Offline mode</span>'),
                unsafe_allow_html=True)
    left, right = st.columns([1.15, 1], gap="large")

    # ---------- LEFT: Details + Knowledge ----------
    with left:
        with st.container(border=True):
            st.markdown("#### 🧩 Details")
            st.text_input("Name", key="agent_name")
            st.text_area("Description", key="agent_description", height=68)
            st.markdown("**Select your agent's model**")
            st.selectbox("Model", ["Claude", "Groq"], index=0, key="llm_provider",
                         label_visibility="collapsed",
                         help="Falls back to the other provider, then offline, if a key is missing.")
            st.caption(f"Active backend: `{backend}`  ·  "
                       + ("🟢 live LLM" if online else "⚪ offline"))
            st.markdown("**Instructions**")
            st.text_area("Instructions", key="agent_instructions", height=96,
                         label_visibility="collapsed",
                         placeholder="Describe what the agent should do, its tone and rules. "
                                     "These guide every answer.")

        with st.container(border=True):
            st.markdown("#### 📚 Knowledge")
            st.caption("Add data, files and resources to inform answers. "
                       "PDFs & images are scanned for text, tables and diagrams.")
            files = st.file_uploader(
                "Add knowledge",
                type=["pdf", "docx", "csv", "xlsx", "pptx", "txt", "md", "png", "jpg", "jpeg"],
                accept_multiple_files=True, label_visibility="collapsed")
            c1, c2 = st.columns([1, 1])
            scan = c1.checkbox("🖼️ Scan diagrams & charts", value=False,
                               help="Azure Document Intelligence (text/tables) + LLM vision "
                                    "for architecture diagrams/graphs. Needs keys to activate.")
            if c2.button("🔍 Add to knowledge", use_container_width=True):
                if files:
                    with st.spinner("Processing…"):
                        n = ingest_files(files, scan)
                    st.success(f"Added {n} file(s). Indexed with "
                               f"`{st.session_state.retriever_backend}` search.")
                else:
                    st.info("Choose file(s) first.")
            docs = st.session_state.documents
            if docs:
                st.markdown(f"**{len(docs)} source(s)** · index "
                            f"`{st.session_state.retriever_backend}`")
                # Highlighted analysis of the most recently added document.
                latest = docs[-1]
                with st.container(border=True):
                    st.markdown(f"##### 📋 Analysis · {latest['name']}")
                    an = latest.get("analysis", {})
                    for sec in ("Objective", "Current solutions",
                                "Proposed solutions", "Insights"):
                        val = (an.get(sec) or "").strip()
                        st.markdown(f"**{sec}**")
                        if val:
                            for ln in val.splitlines():
                                st.markdown(f"- {ln}" if not ln.startswith("-") else ln)
                        else:
                            st.caption("—")
                for i, d in enumerate(docs):
                    with st.expander(f"📄 {d['name']}  ·  {d['domain']}  ·  {d['chars']} chars"):
                        an = d.get("analysis", {})
                        for sec in ("Objective", "Current solutions",
                                    "Proposed solutions", "Insights"):
                            v = (an.get(sec) or "").strip()
                            if v:
                                st.markdown(f"**{sec}:** {v}")
                        vis = d.get("visuals", {})
                        if vis.get("descriptions"):
                            st.markdown("**Diagrams / visuals:**")
                            for v in vis["descriptions"]:
                                st.caption(f"p{v.get('page','?')}: {v.get('description','')[:400]}")
                        elif vis.get("available") is False and vis.get("note"):
                            st.caption(vis["note"])
                        if st.button("Remove", key=f"rm_{i}"):
                            audit("KNOWLEDGE_REMOVE", d["name"])
                            st.session_state.documents.pop(i)
                            rebuild_index(); st.rerun()
            else:
                st.info("No knowledge yet — add a document to start.")

    # ---------- RIGHT: Test your agent (chat) ----------
    with right:
        with st.container(border=True):
            st.markdown("#### 💬 Test your agent")
            if not online:
                st.warning("No online LLM connected — answers are offline extraction. "
                           "Set **CLAUDE_API_KEY** or **GROQ_API_KEY** for live answers.")
            with st.chat_message("assistant"):
                st.write(f"Hello, I'm {st.session_state.agent_name}. "
                         "How can I help you today?")
            for m in st.session_state.history:
                with st.chat_message(m["role"]):
                    st.write(m["content"])
                    if m["role"] == "assistant" and not m.get("blocked"):
                        meta = []
                        if m.get("domain") and m["domain"] != "General":
                            meta.append(f'<span class="pf-dom">{m["domain"]}</span>')
                        st.markdown(" ".join(meta), unsafe_allow_html=True)
                        if m.get("sources"):
                            st.caption("📎 " + ", ".join(m["sources"]))
                        if m.get("confidence") is not None:
                            st.caption(f"confidence {m['confidence']:.0%} · "
                                       f"grounded {m.get('grounded',0):.0%}")

            # download button appears once a report has been generated in chat
            if st.session_state.get("report_file"):
                fname, data = st.session_state["report_file"]
                st.download_button(
                    "⬇️ Download proposed-solutions report (.docx)", data=data,
                    file_name=fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True)

            # follow-ups based on the PREVIOUS question
            last_q = next((m["content"] for m in reversed(st.session_state.history)
                           if m["role"] == "user"), "")
            last_a = next((m["content"] for m in reversed(st.session_state.history)
                           if m["role"] == "assistant"), "")
            if last_q:
                fups = followups(last_q, last_a)
                if fups:
                    st.markdown("**💡 Follow-up questions**")
                    fcols = st.columns(len(fups))
                    for i, f in enumerate(fups):
                        if fcols[i].button(f"❓ {f[:48]}", key=f"f_{len(st.session_state.history)}_{i}"):
                            with st.spinner("🤖 Thinking…"):
                                answer_query(f)
                            st.rerun()

            q = st.chat_input("Ask a question or describe what you need")
            if q:
                with st.spinner("🤖 Thinking…"):
                    answer_query(q)
                st.rerun()

# =========================== AGENTS (LangGraph) ============================ #
with tab_agents:
    st.markdown("### 🤖 Agents — LangGraph pipeline")
    st.caption("Runs Route → Guardrails → Retrieve → Generate → Suggest over ALL your "
               "uploaded documents (multi-doc, multi-domain), with latency, token usage "
               "and an LLM-graded RAGAS score.")
    mode = st.radio(
        "Execution mode", ["LangGraph Engine (real graph over your docs)",
                           "Instrumented pipeline (per-stage timing)"],
        horizontal=True, label_visibility="collapsed")
    engine_mode = mode.startswith("LangGraph Engine")
    if not st.session_state.documents:
        st.info("Add documents in the Overview tab first — the pipeline runs on your "
                "uploaded knowledge across domains.")
    if not online:
        st.info("Offline mode — the pipeline still runs and is timed; set "
                "CLAUDE_API_KEY / GROQ_API_KEY for live answers and exact token counts.")
    if engine_mode:
        n = feed_engine_docs()
        eng = get_engine()
        try:
            stats = eng.store.stats()
            ns = {k: v for k, v in stats.items()}
            st.caption("Engine knowledge namespaces (domain → chunks): " +
                       " · ".join(f"{k}:{v}" for k, v in ns.items()))
        except Exception:
            pass

    aq = st.text_input("Ask the agent pipeline",
                       placeholder="e.g. Across all documents, what are the top problems and how do we solve them?")
    if st.button("▶️ Run pipeline") and aq:
        if engine_mode:
            # ---- REAL LangGraph Engine graph over the fed documents ----
            feed_engine_docs()
            eng = get_engine()
            t = time.perf_counter()
            try:
                turn = eng.ask_graph(aq)
                if getattr(turn, "status", "") == "ask_domain":
                    # single-shot run: confirm the router's suggested domain (or the
                    # query's classified domain) so the graph proceeds to an answer
                    cd = getattr(turn, "domain", None) or \
                        ENGINE_DOMAIN_MAP.get(classify_domain(aq), "general")
                    turn = eng.ask_graph(aq, confirmed_domain=cd)
            except Exception as ex:
                turn = None; _err = str(ex)
            total = (time.perf_counter() - t) * 1000
            if turn is not None:
                a = getattr(turn, "answer", None)
                ans_text = a.text if a else (getattr(turn, "message", "") or "(no answer)")
                dom = getattr(turn, "domain", None) or classify_domain(aq)
                fups = list(getattr(turn, "follow_ups", []) or [])
            else:
                ans_text, dom, fups = f"(engine error: {_err})", classify_domain(aq), []
            # context the graph retrieved (same store.retrieve) for RAGAS + evidence
            ctx_chunks, evidence = [], []
            try:
                for ch, _score in eng.store.retrieve(aq, dom):
                    ctx_chunks.append(ch.text)
                    src = getattr(ch, "source", None) or getattr(ch, "section", "?")
                    if src not in [e["source"] for e in evidence]:
                        evidence.append({"source": src, "domain": dom})
            except Exception:
                pass
            ctx = "\n\n".join(ctx_chunks)
            tokens = int(getattr(eng.llm, "last_tokens", 0) or 0) or (len(aq) + len(ans_text)) // 4
            stages = None
        else:
            # ---- Instrumented per-stage pipeline over the UI retriever ----
            stages, tokens = {}, 0
            t = time.perf_counter(); dom = classify_domain(aq)
            stages["route"] = (time.perf_counter() - t) * 1000
            t = time.perf_counter(); _ = guard.check_pii(aq)
            stages["guard"] = (time.perf_counter() - t) * 1000
            t = time.perf_counter()
            r = st.session_state.retriever; ret = []
            if r is not None:
                try: ret = r.search(aq, top_k=5)
                except Exception: ret = []
            stages["retrieve"] = (time.perf_counter() - t) * 1000
            ctx = "\n\n".join(f"[Source: {h['source']}]\n{h['text']}" for h in ret)
            dom_by_name = {d["name"]: d.get("domain", "General")
                           for d in st.session_state.documents}
            evidence = []
            for h in ret:
                nm = h.get("source", "?")
                if nm not in [e["source"] for e in evidence]:
                    evidence.append({"source": nm, "domain": dom_by_name.get(nm, "General")})
            t = time.perf_counter(); ans_text = ""
            prompt = (f"Documents (sources may span domains):\n{ctx}\n\nQuestion: {aq}\n"
                      "Answer directly, cite which document/domain each point draws from, "
                      "and give proposed solutions and insights.")
            if online and hasattr(client, "complete"):
                try:
                    ans_text = client.complete(SYSTEM_SOLUTIONS, prompt) or ""
                    tokens = int(getattr(client, "last_tokens", 0) or 0)
                except Exception:
                    ans_text = ""
            if not ans_text:
                resp = client.answer_question(ctx, aq,
                    allow_general_knowledge=not bool(st.session_state.documents))
                ans_text = resp.get("answer", "") or "(no answer)"
                tokens = int(resp.get("tokens", 0) or 0)
            if not tokens:
                tokens = (len(prompt) + len(ans_text)) // 4
            stages["generate"] = (time.perf_counter() - t) * 1000
            t = time.perf_counter(); fups = followups(aq, ans_text)
            stages["suggest"] = (time.perf_counter() - t) * 1000

        total = sum(stages.values()) if stages else total
        st.session_state.setdefault("_agent_lat", []).append(round(total, 1))
        st.session_state["_agent_lat"] = st.session_state["_agent_lat"][-50:]
        st.session_state.setdefault("_agent_tok", []).append(int(tokens))
        st.session_state["_agent_tok"] = st.session_state["_agent_tok"][-50:]
        st.session_state["_agent_last"] = {"dom": dom, "ans": ans_text, "fups": fups,
                                           "stages": stages, "total": total, "ctx": ctx,
                                           "aq": aq, "tokens": tokens, "evidence": evidence,
                                           "engine": engine_mode}
        audit("AGENT_PIPELINE", f"{'engine' if engine_mode else 'instrumented'} · {dom} · "
              f"{total:.0f}ms · {tokens} tok · {len(evidence)} sources")

    last = st.session_state.get("_agent_last")
    if last:
        tag = "LangGraph Engine" if last.get("engine") else "Instrumented"
        st.markdown(f'**Mode:** {tag} · **Domain:** '
                    f'<span class="pf-dom">{last["dom"]}</span>', unsafe_allow_html=True)
        st.write(last["ans"])
        if last["evidence"]:
            st.markdown("**Evidence used (multi-doc / multi-domain):** " +
                        " ".join(f'<span class="pf-dom">{e["source"]} · {e["domain"]}</span>'
                                 for e in last["evidence"]), unsafe_allow_html=True)
        if last["fups"]:
            st.markdown("**Follow-ups:** " + " · ".join(last["fups"]))
        st.divider()
        st.markdown("#### ⏱️ Latency & tokens (this run)")
        if last.get("stages"):
            cols = st.columns(5)
            for col, k in zip(cols, ["route", "guard", "retrieve", "generate", "suggest"]):
                col.metric(k.title(), f"{last['stages'].get(k, 0):.0f} ms")
        t1, t2 = st.columns(2)
        t1.metric("Total pipeline latency", f"{last['total']:.0f} ms")
        t2.metric("Tokens used", f"{last['tokens']:,}")
        if last.get("stages"):
            st.bar_chart({k: round(v, 1) for k, v in last["stages"].items()})
        st.markdown("#### 🧪 RAGAS evaluation (LLM-graded)")
        rag = evaluate_metrics(last["aq"], last["ans"],
                               [last["ctx"]] if last["ctx"] else [],
                               client=client, online=online)
        st.caption("Graded by: " + rag.get("mode", "?"))
        rc = st.columns(5)
        rc[0].metric("Faithfulness", f"{rag['faithfulness']:.0%}")
        rc[1].metric("Answer relevance", f"{rag['answer_relevance']:.0%}")
        rc[2].metric("Context precision", f"{rag['context_precision']:.0%}")
        rc[3].metric("Context recall", f"{rag['context_recall']:.0%}")
        rc[4].metric("Overall", f"{rag['overall']:.0%}")

    lats = st.session_state.get("_agent_lat", [])
    toks = st.session_state.get("_agent_tok", [])
    if lats:
        import statistics as _stx
        st.divider(); st.markdown("#### 📈 Production metrics")
        p50 = _stx.median(lats)
        p95 = sorted(lats)[max(0, int(round(len(lats) * 0.95)) - 1)]
        g = st.columns(4)
        g[0].metric("Runs", len(lats))
        g[1].metric("p50 latency", f"{p50:.0f} ms")
        g[2].metric("p95 latency", f"{p95:.0f} ms")
        g[3].metric("Total tokens", f"{sum(toks):,}")
        gc1, gc2 = st.columns(2)
        with gc1:
            st.markdown("**Latency per run (ms)**"); st.line_chart({"latency_ms": lats})
        with gc2:
            st.markdown("**Tokens per run**"); st.bar_chart({"tokens": toks})


# =========================== COMPLIANCE MATRIX ============================ #
with tab_compliance:
    st.markdown("### 📋 Compliance / Requirements Traceability")
    st.caption("Extract requirements from an RFP and trace each to your proposal "
               "response — Covered / Partial / Missing. "
               "Built via Kiro SDD (.kiro/specs/compliance-matrix/).")

    docs = st.session_state.documents
    prefill_rfp = docs[0]["text"][:4000] if docs else ""
    prefill_resp = docs[-1]["text"][:4000] if len(docs) > 1 else ""
    cc1, cc2 = st.columns(2)
    rfp_text = cc1.text_area("RFP / requirements source", value=prefill_rfp,
                             height=220, placeholder="Paste the RFP text (shall/must/…)")
    resp_text = cc2.text_area("Our proposal response", value=prefill_resp,
                              height=220, placeholder="Paste our proposal response")
    if docs:
        st.caption("Tip: pre-filled from your uploaded Knowledge — edit as needed.")

    if st.button("📋 Build traceability matrix"):
        if not rfp_text.strip():
            st.info("Paste some RFP text first.")
        else:
            m = build_matrix(rfp_text, resp_text, llm=client)
            s = m.summary
            k1, k2, k3, k4 = st.columns(4)
            k1.metric("Compliance", f"{s['compliance_pct']}%")
            k2.metric("Covered", s["covered"])
            k3.metric("Partial", s["partial"])
            k4.metric("Missing", s["missing"])
            badge = {"Covered": "🟢", "Partial": "🟠", "Missing": "🔴"}
            st.dataframe(
                [{"ID": r.id, "Status": f"{badge[r.status]} {r.status}",
                  "Requirement": r.requirement,
                  "Match": f"{r.score:.0%}", "Evidence": r.evidence[:120]}
                 for r in m.rows],
                use_container_width=True, hide_index=True)
            if m.gaps:
                st.markdown("**⚠️ Gaps to address (worst first)**")
                for g in m.gaps:
                    st.markdown(f"- `{g.id}` **{g.status}** — {g.requirement}")
            st.caption("🧠 " + m.gap_summary)
            audit("COMPLIANCE_MATRIX",
                  f"{s['compliance_pct']}% · {s['covered']}/{s['total']} covered")

# =========================== EVALUATION (RAGAS) ============================ #
with tab_eval:
    st.markdown("### 🧪 Evaluation — RAGAS (LLM-graded)")
    st.caption("Each question is answered from the retrieved context, then graded on "
               "four RAGAS metrics — faithfulness, answer relevance, context precision "
               "and context recall. With a key set, metrics are LLM-graded; offline it "
               "falls back to a lexical heuristic.")
    default_qs = ("What is the main objective?\nWhat are the key problems?\n"
                  "What solutions are proposed?")
    qs = st.text_area("Questions (one per line)", value=default_qs, height=90)
    if st.button("▶️ Run RAGAS evaluation"):
        if not st.session_state.documents:
            st.info("Add knowledge in the Overview tab first.")
        else:
            rows, agg, modes = [], {k: [] for k in METRIC_KEYS + ["overall"]}, set()
            r = st.session_state.retriever
            prog = st.progress(0.0)
            questions = [x.strip() for x in qs.splitlines() if x.strip()]
            for i, q in enumerate(questions):
                ret = []
                if r is not None:
                    try: ret = r.search(q, top_k=4)
                    except Exception: ret = []
                chunks = [h["text"] for h in ret] or \
                    [st.session_state.documents[-1]["text"][:3000]]
                ctx = "\n\n".join(chunks)
                a = ""
                if online and hasattr(client, "complete"):
                    try:
                        a = client.complete(SYSTEM_SOLUTIONS,
                            f"Documents:\n{ctx}\n\nQuestion: {q}\n"
                            "Answer directly and grounded in the documents.") or ""
                    except Exception:
                        a = ""
                if not a:
                    a = client.answer_question(ctx, q,
                                               allow_general_knowledge=False).get("answer", "")
                m = evaluate_metrics(q, a, chunks, client=client, online=online)
                modes.add(m.get("mode", "?"))
                for k in agg:
                    agg[k].append(m[k])
                rows.append({"Question": q,
                             "Faithfulness": f"{m['faithfulness']:.0%}",
                             "Answer relevance": f"{m['answer_relevance']:.0%}",
                             "Context precision": f"{m['context_precision']:.0%}",
                             "Context recall": f"{m['context_recall']:.0%}",
                             "Overall": f"{m['overall']:.0%}"})
                prog.progress((i + 1) / len(questions))
            prog.empty()
            mean = lambda v: (sum(v) / len(v)) if v else 0.0
            st.caption("Graded by: " + " · ".join(sorted(modes)))
            m = st.columns(5)
            m[0].metric("Faithfulness", f"{mean(agg['faithfulness']):.0%}")
            m[1].metric("Answer relevance", f"{mean(agg['answer_relevance']):.0%}")
            m[2].metric("Context precision", f"{mean(agg['context_precision']):.0%}")
            m[3].metric("Context recall", f"{mean(agg['context_recall']):.0%}")
            m[4].metric("Overall", f"{mean(agg['overall']):.0%}")
            st.dataframe(rows, use_container_width=True, hide_index=True)
            st.markdown("**RAGAS metric averages**")
            st.bar_chart({k.replace('_', ' ').title(): round(mean(agg[k]), 3)
                          for k in METRIC_KEYS})
            audit("EVALUATION",
                  f"RAGAS · {len(rows)} q · overall {mean(agg['overall']):.0%}")

# =========================== ANALYTICS ===================================== #
with tab_analytics:
    st.markdown("### 📊 Analytics")
    hist = st.session_state.history
    asks = [m for m in hist if m["role"] == "user"]
    ans = [m for m in hist if m["role"] == "assistant" and not m.get("blocked")]
    confs = [m["confidence"] for m in ans if m.get("confidence") is not None]
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("📚 Documents", len(st.session_state.documents))
    c2.metric("💬 Questions", len(asks))
    c3.metric("🎯 Avg confidence", f"{(sum(confs)/len(confs)):.0%}" if confs else "—")
    c4.metric("🛡️ Guardrail hits", st.session_state.guardrail_hits)
    doms = {}
    for d in st.session_state.documents:
        doms[d["domain"]] = doms.get(d["domain"], 0) + 1
    if doms:
        st.markdown("**Knowledge by domain**")
        st.bar_chart(doms)
    if confs:
        st.markdown("**Answer confidence over time**")
        st.line_chart({"confidence": confs})

# =========================== ACTIVITY (audit log) ========================= #
with tab_activity:
    st.markdown("### 🧾 Activity — audit log")
    tracker = get_tracker()
    st.caption(f"Persisted to **{tracker.describe()}**. In production this is Azure "
               "SQL (set AZURE_SQL_CONNECTION_STRING); locally it uses SQLite.")
    rows = tracker.get_logs(limit=300) or st.session_state.audit
    if rows:
        by = {}
        for a in rows:
            by[a["action"]] = by.get(a["action"], 0) + 1
        cc = st.columns(min(len(by), 6) or 1)
        for i, (k, v) in enumerate(list(by.items())[:6]):
            cc[i % len(cc)].metric(str(k).replace("_", " ").title(), v)
        st.dataframe(rows, use_container_width=True, hide_index=True)
        st.download_button("⬇️ Download audit log (CSV)", data=_audit_csv(rows),
                           file_name="audit_log.csv", mime="text/csv")
    else:
        st.info("No activity yet — upload knowledge or ask a question in Overview.")

# =========================== TEST SUITES =================================== #
with tab_tests:
    st.markdown("### ✅ Test Suites")
    st.caption("Unit + integration tests (pytest) and an end-to-end load test — "
               "run them here and see pass/fail, coverage and latency.")

    st.markdown("#### 🧪 Unit & integration tests + coverage")
    if st.button("▶️ Run test suite"):
        with st.spinner("Running pytest…"):
            st.session_state._tests = run_test_suite(str(Path(__file__).resolve().parents[2]))
    tr = st.session_state.get("_tests")
    if tr and tr.get("ok"):
        t = tr["tests"]["totals"]
        a, b, c, d = st.columns(4)
        a.metric("Total", t["total"]); b.metric("Passed", t["passed"])
        c.metric("Failed", t["failed"] + t["errors"])
        d.metric("Coverage", f"{tr.get('coverage',{}).get('total_percent','?')}%")
    elif tr:
        st.error(tr.get("error", "Test run failed"))

    st.divider()
    st.markdown("#### 🚀 Load test (end-to-end)")
    st.caption("Phased per-endpoint latency & sustained RPS vs budgets "
               "(health <50ms, retrieve <150ms, end-to-end p95 <5s).")
    lc1, lc2, lc3 = st.columns(3)
    base = lc1.text_input("Target API base URL",
                          value=os.getenv("SELF_API_BASE", "http://127.0.0.1:8001"))
    users = lc2.slider("Concurrent users", 1, 50, 10)
    dur = lc3.slider("Duration (s)", 3, 60, 8)
    if st.button("▶️ Run load test"):
        with st.spinner(f"Load testing {base}…"):
            st.session_state._load = run_load_test(base, users, dur)
    lr = st.session_state.get("_load")
    if lr and lr.get("ok"):
        agg = lr["aggregate"]
        a, b, c = st.columns(3)
        a.metric("🔥 Peak RPS", agg["peak_rps"]); b.metric("Requests", agg["total_requests"])
        c.metric("Result", "PASS ✅" if agg["all_pass"] else "FAIL ❌")
        st.dataframe([{"Endpoint": r["endpoint"], "RPS": r["rps"], "p50 (ms)": r["p50_ms"],
                       "p95 (ms)": r["p95_ms"], "Budget (ms)": r["budget_ms"],
                       "Status": r["status"]} for r in lr["endpoints"]],
                     use_container_width=True, hide_index=True)

st.markdown("<hr><p style='text-align:center;color:#5C6B82'>"
            "<b>ProposalForge Agent</b> · LangGraph agentic RAG · Claude/Groq · Pinecone · "
            "Azure Monitor</p>", unsafe_allow_html=True)
