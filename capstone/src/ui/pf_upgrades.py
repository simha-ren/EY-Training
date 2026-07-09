"""ProposalForge UI upgrades — in-chat charts + off-topic relevance gate.

Drop-in module used by src/ui/app_prod.py. No other files need to change beyond
three small hooks in app_prod.py (see MODIFICATIONS.md).

Design goals:
  * Works ONLINE (LLM emits grounded chart JSON) and OFFLINE (parses numeric
    series out of the retrieved context / CSV-like tables) — so a live demo
    never shows an empty chart.
  * Off-topic detection is conservative: it only fires when documents are loaded
    AND the question has no lexical grounding in them, so genuine document
    questions are never wrongly blocked.
"""
from __future__ import annotations

import os
import re
import json
from typing import Optional, List, Dict


# --------------------------------------------------------------------------- #
#  Chart intent + extraction                                                  #
# --------------------------------------------------------------------------- #
_CHART_WORDS = re.compile(
    r"\b(chart|graph|plot|trend|trends|visuali[sz]e|visuali[sz]ation|bar chart|"
    r"line chart|over (the )?(last |past |next )?\d+\s*(year|years|quarter|quarters|month|months)|"
    r"year[- ]on[- ]year|growth curve|show me the (numbers|data))\b", re.I)


def chart_intent(query: str) -> bool:
    """True when the user is asking to SEE data as a chart."""
    return bool(_CHART_WORDS.search(query or ""))


_STOP = set("the a an of to in for and or is are on with this that what which how why "
            "when who from as at by be it its into show me give the numbers data".split())


def _tok(s: str) -> List[str]:
    return [w for w in re.findall(r"[a-z0-9]{3,}", (s or "").lower()) if w not in _STOP]


def _numbers_from_context(context: str, query: str) -> Optional[Dict]:
    """Offline fallback — pull a labelled numeric series out of CSV/table text.

    Recognises lines like  '2021, 95, 310'  or  'FY22 | 180 | 610'  and years/
    quarters as x-axis labels. Returns a chart spec or None.
    """
    rows = []
    for ln in (context or "").splitlines():
        parts = re.split(r"[|,\t]| {2,}", ln.strip())
        parts = [p.strip() for p in parts if p.strip()]
        if len(parts) < 2:
            continue
        label = parts[0]
        nums = []
        for p in parts[1:]:
            m = re.search(r"-?\d[\d,\.]*", p.replace("%", "").replace("INR", ""))
            if m:
                try:
                    nums.append(float(m.group(0).replace(",", "")))
                except ValueError:
                    pass
        if nums and re.search(r"(19|20)\d{2}|FY\d{2}|Q\d|Year\s*\d|Wk\d", label, re.I):
            rows.append((label, nums))
    if len(rows) < 2:
        return None
    width = min(len(r[1]) for r in rows)
    x = [r[0] for r in rows]
    series = {f"series {i+1}": [r[1][i] for r in rows] for i in range(min(width, 3))}
    return {"kind": "line", "title": "Trend from your documents", "x": x, "series": series}


def build_chart(query: str, context: str, client, online: bool) -> Optional[Dict]:
    """Return a chart spec {kind,title,x,series} grounded in the documents, or None."""
    # 1) ONLINE — ask the model for grounded chart JSON.
    if online and hasattr(client, "complete"):
        try:
            raw = client.complete(
                "You extract chart-ready data from documents. Reply with ONLY strict "
                "JSON, no prose, no code fences.",
                "From the document context, produce data to answer the user's chart "
                "request. Use ONLY numbers present or clearly implied in the context.\n"
                'JSON schema: {"kind":"line|bar","title":"...","x":["label",...],'
                '"series":{"Series name":[num,...]}}\n'
                "If the context has no chartable numbers, reply exactly: {}\n\n"
                f"User request: {query}\n\nContext:\n{context[:5000]}")
            raw = re.sub(r"```json|```", "", raw or "").strip()
            a, b = raw.find("{"), raw.rfind("}")
            if a >= 0 and b > a:
                spec = json.loads(raw[a:b + 1])
                if spec.get("x") and spec.get("series"):
                    spec.setdefault("kind", "line")
                    spec.setdefault("title", "Chart")
                    return spec
        except Exception:
            pass
    # 2) OFFLINE — parse numbers out of the retrieved context.
    return _numbers_from_context(context, query)


def render_chart(st, spec: Dict) -> None:
    """Render a chart spec inside a Streamlit chat message."""
    try:
        import pandas as pd
    except Exception:
        st.caption("(install pandas to render charts)")
        return
    x = spec.get("x") or []
    series = spec.get("series") or {}
    if not x or not series:
        return
    n = len(x)
    data = {k: (list(v) + [None] * n)[:n] for k, v in series.items()}
    df = pd.DataFrame(data, index=x)
    st.markdown(f"**📈 {spec.get('title','Chart')}**")
    if (spec.get("kind") or "line").lower().startswith("bar"):
        st.bar_chart(df)
    else:
        st.line_chart(df)


# --------------------------------------------------------------------------- #
#  Off-topic relevance gate                                                   #
# --------------------------------------------------------------------------- #
def relevance(query: str, context: str) -> float:
    """Lexical overlap between the query and the retrieved context (0..1)."""
    q = set(_tok(query))
    c = set(_tok(context))
    if not q:
        return 1.0
    return len(q & c) / len(q)


# Questions ABOUT the documents (summaries, problems, solutions, insights, etc.)
# are never off-topic when documents are loaded — even though they share no
# content words with the corpus (e.g. "summarise this doc" scores 0 in TF-IDF).
_DOC_META = re.compile(
    r"\b(summar|overview|objective|goal|challenge|problem|issue|risk|gap|"
    r"solution|propos|recommend|insight|takeaway|finding|highlight|key point|"
    r"main point|current situation|address|improve|mitigat|fix|analy|explain|describe|"
    r"what does|what is in|whats in|what's in|tell me about|about (this|these|the) "
    r"(doc|document|file|report|proposal)|across (all )?(doc|domain)|compare|"
    r"table|figure|chart|graph|diagram|budget|kpi|metric|timeline)", re.I)

# TF-IDF cosine below this means the query shares no meaningful terms with ANY
# document — i.e. genuinely off-topic. Off-topic queries score ~0.0; real
# content queries score >0.1 (env-tunable).
_OFFTOPIC_MIN_SCORE = float(os.getenv("OFFTOPIC_MIN_SCORE", "0.03"))


def offtopic_banner(query, context, docs, retrieved, domains=None):
    """Return a heads-up banner ONLY when a question is genuinely unrelated to the
    uploaded documents. Conservative by design — document questions (including
    'summarise this', 'what are the challenges', 'proposed solutions') are never
    flagged. Fires for things like 'what is the capital of India' when docs exist.
    """
    if not docs:
        return None                                   # nothing to be off-topic from
    if _DOC_META.search(query or ""):
        return None                                   # it's a question about the docs
    # Use the retriever's own top similarity score as the signal.
    top = 0.0
    for h in (retrieved or []):
        try:
            top = max(top, float(h.get("score", 0.0)))
        except (TypeError, ValueError):
            pass
    if top >= _OFFTOPIC_MIN_SCORE:
        return None                                   # matches the corpus — on topic
    doms = ", ".join(sorted({d.get("domain", "General") for d in docs})) or "your documents"
    return (f"🧭 **Heads up — this looks outside your uploaded knowledge** "
            f"({doms}). I couldn't find it in your documents, so the answer below "
            f"is general knowledge, not grounded in your sources.")


# --------------------------------------------------------------------------- #
#  Downloadable report — embedded chart (graphs/visualizations requirement)   #
# --------------------------------------------------------------------------- #
def _count(analysis: dict, key: str) -> int:
    v = (analysis or {}).get(key, "") or ""
    if isinstance(v, list):
        return len([x for x in v if str(x).strip()])
    return len([ln for ln in str(v).splitlines() if ln.strip()])


def add_report_chart(document, docs: list) -> None:
    """Embed a 'challenges vs proposed solutions per document' chart into a docx.

    Uses matplotlib when installed (crisp PNG); otherwise falls back to a
    dependency-free shaded-table bar chart so the report always has a visual.
    """
    if not docs:
        return
    names = [(d.get("name", "doc")[:24]) for d in docs]
    challenges = [_count(d.get("analysis", {}), "Current solutions") for d in docs]
    solutions = [_count(d.get("analysis", {}), "Proposed solutions") for d in docs]

    # --- rich path: matplotlib PNG ---
    try:
        import io
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from docx.shared import Inches
        x = range(len(names))
        fig, ax = plt.subplots(figsize=(6.6, 3.4))
        ax.bar([i - 0.2 for i in x], challenges, width=0.4, label="Current challenges", color="#E9950C")
        ax.bar([i + 0.2 for i in x], solutions, width=0.4, label="Proposed solutions", color="#12A5B0")
        ax.set_xticks(list(x))
        ax.set_xticklabels(names, rotation=25, ha="right", fontsize=8)
        ax.set_title("Challenges vs Proposed Solutions by Document", fontsize=11, fontweight="bold")
        ax.legend(frameon=False, fontsize=8)
        ax.grid(axis="y", color="#eef2f6")
        fig.tight_layout()
        buf = io.BytesIO(); fig.savefig(buf, format="png", dpi=130); plt.close(fig)
        buf.seek(0)
        document.add_heading("Portfolio at a glance", level=1)
        document.add_picture(buf, width=Inches(6.0))
        return
    except Exception:
        pass

    # --- fallback path: shaded-table bar chart (no extra deps) ---
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    document.add_heading("Portfolio at a glance", level=1)
    document.add_paragraph("Current challenges (amber) vs proposed solutions (teal) per document:")
    maxv = max([1] + challenges + solutions)

    def shade(cell, hexcolor):
        tcPr = cell._tc.get_or_add_tcPr()
        sh = OxmlElement("w:shd"); sh.set(qn("w:fill"), hexcolor); tcPr.append(sh)

    tbl = document.add_table(rows=0, cols=2)
    for nm, ch, so in zip(names, challenges, solutions):
        for label, val, color in ((f"{nm} — challenges", ch, "E9950C"),
                                   (f"{nm} — solutions", so, "12A5B0")):
            row = tbl.add_row().cells
            row[0].text = f"{label} ({val})"
            bar = row[1]
            blocks = int(round((val / maxv) * 20))
            bar.text = "█" * max(blocks, 1)
            shade(bar, color)


# --------------------------------------------------------------------------- #
#  Fast, contextual follow-up suggestions (instant, no LLM round-trip needed)  #
# --------------------------------------------------------------------------- #
_ASK_CHALLENGE = re.compile(r"challenge|problem|issue|risk|current situation|gap|weak", re.I)
_ASK_SOLUTION = re.compile(r"solution|propos|recommend|fix|solve|improve|address|how (to|do|can)", re.I)
_ASK_INSIGHT = re.compile(r"insight|takeaway|finding|so what|implication|highlight", re.I)
_ASK_SUMMARY = re.compile(r"summar|overview|objective|goal|what is (this|in)|about (this|the)|describe|explain", re.I)
_ASK_CHART = re.compile(r"chart|graph|plot|trend|over \d+ year|growth|compare|number|metric", re.I)


def smart_followups(last_q: str, last_a: str, docs: list, offtopic: bool = False):
    """Return up to 3 relevant, clickable follow-up questions for the last turn.

    Instant (no API call) and always non-empty, so suggestions never disappear in
    a demo. Style matches the product spec: 'Do you need the proposed solutions?',
    'Do you want insights on this?', plus one topical/related question.
    """
    q = (last_q or "").strip()
    doms = [d.get("domain", "") for d in (docs or []) if d.get("domain")]
    dom = next((d for d in doms if d and d != "General"), "these documents")

    if offtopic:
        return ["Summarise the uploaded documents",
                "What are the current challenges?",
                "What are the proposed solutions?"]

    out: list = []
    def add(s):
        if s and s not in out and len(out) < 3:
            out.append(s)

    if _ASK_CHALLENGE.search(q):
        add("Do you want the proposed solutions to these challenges?")
        add("Any insights I should check on for these?")
        add("Show me a chart of the key metrics")
    elif _ASK_SOLUTION.search(q):
        add("What current challenges do these solutions address?")
        add("What insights support these recommendations?")
        add("Summarise the objective and expected impact")
    elif _ASK_INSIGHT.search(q):
        add("Do you need the proposed solutions behind these insights?")
        add("What are the current challenges here?")
        add("Show these insights as a chart")
    elif _ASK_SUMMARY.search(q):
        add("Do you need the problems in these docs?")
        add("Do you need the objectives?")
        add("Do you need the proposed solutions to it?")
    elif _ASK_CHART.search(q):
        add("Do you want the proposed solutions for this trend?")
        add("What are the key insights from this data?")
        add("Summarise what this means for the objective")
    else:
        # generic but still on the product's rails — the requested phrasing
        add("Do you need the problems in these docs?")
        add("Do you need the objectives?")
        add("Do you need the proposed solutions to it?")

    # always leave the user a strong next step
    add("Compare the challenges across all documents")
    return out[:3]


def followups_fast(client, last_q: str, last_a: str, docs: list,
                   guard=None, offtopic: bool = False, use_llm: bool = False):
    """Fast path for the UI: heuristic suggestions first (instant + relevant); if
    use_llm is set and online, try to enrich with model suggestions, but never
    block or return empty. PII-filter the results when a guard is supplied.
    """
    sugg = []
    if use_llm and client is not None and hasattr(client, "get_auto_suggestions"):
        try:
            sugg = list(client.get_auto_suggestions((last_a or "")[:1500], last_q or "") or [])
        except Exception:
            sugg = []
    if not sugg:
        sugg = smart_followups(last_q, last_a, docs, offtopic=offtopic)
    if guard is not None:
        try:
            sugg = [s for s in sugg if not guard.check_pii(s).triggered]
        except Exception:
            pass
    # de-dup, keep 3
    seen, out = set(), []
    for s in sugg:
        if s and s not in seen:
            seen.add(s); out.append(s)
    return out[:3] or smart_followups(last_q, last_a, docs, offtopic=offtopic)


# --------------------------------------------------------------------------- #
#  Bulletize — turn a section blob into clean point-wise bullets               #
# --------------------------------------------------------------------------- #
def bulletize(text, max_points: int = 8):
    """Split an analysis section (which may be a paragraph or newline blob) into
    clean, point-wise bullets. Handles existing '-', '*', '•', numbered lists,
    and long paragraphs (splits on sentences)."""
    if isinstance(text, (list, tuple)):
        lines = [str(x) for x in text]
    else:
        lines = str(text or "").splitlines()
    points = []
    for ln in lines:
        s = ln.strip()
        if not s:
            continue
        s = re.sub(r"^\s*(?:[-*•]|\d+[.)])\s*", "", s).strip()  # strip bullet/number
        # split a multi-sentence blob into separate points
        sentence_count = len(re.findall(r"[.!?](?:\s|$)", s))
        if sentence_count >= 2 or (len(s) > 120 and "." in s):
            for sent in re.split(r"(?<=[.!?])\s+", s):
                sent = sent.strip().rstrip(".")
                if len(sent) > 3:
                    points.append(sent)
        elif len(s) > 2:
            points.append(s.rstrip("."))
    # de-dup preserving order
    seen, out = set(), []
    for p in points:
        k = p.lower()
        if k not in seen:
            seen.add(k); out.append(p)
    return out[:max_points]
