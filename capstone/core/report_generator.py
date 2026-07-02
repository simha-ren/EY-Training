"""Report generation for PDF and DOCX output."""
from __future__ import annotations

from datetime import datetime
from typing import Dict, List, Optional, Any
from pathlib import Path
import json


class ReportGenerator:
    """Generate professional reports in PDF and DOCX formats."""
    
    def __init__(self, output_dir: str = "data/reports"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_analysis_report(self, 
                                document_id: str,
                                filename: str,
                                analysis: Dict[str, Any],
                                conversation_history: List[Dict],
                                insights: Optional[Dict] = None,
                                format: str = "pdf") -> str:
        """Generate comprehensive analysis report."""
        
        # Create report data structure
        report_data = {
            "document_id": document_id,
            "filename": filename,
            "generated_at": datetime.now().isoformat(),
            "objective": analysis.get("objective", ""),
            "challenges": analysis.get("challenges", []),
            "proposed_solutions": analysis.get("proposed_solutions", []),
            "key_insights": analysis.get("insights", []),
            "confidence_level": analysis.get("confidence", 0),
            "conversation_history": conversation_history,
            "additional_insights": insights or {},
        }
        
        if format.lower() == "pdf":
            return self._generate_pdf_report(report_data)
        elif format.lower() in ["docx", "doc"]:
            return self._generate_docx_report(report_data)
        else:
            raise ValueError(f"Unsupported format: {format}")
    
    def _generate_pdf_report(self, report_data: Dict[str, Any]) -> str:
        """Generate PDF report."""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = self.output_dir / f"Report_{timestamp}.pdf"
            
            # Create document
            doc = SimpleDocTemplate(str(filename), pagesize=letter)
            elements = []
            
            # Styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#1f6f6b'),
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.HexColor('#1f6f6b'),
                spaceAfter=12,
                spaceBefore=12
            )
            
            # Title
            elements.append(Paragraph("Analysis Report", title_style))
            elements.append(Spacer(1, 0.2 * inch))
            
            # Document Info
            elements.append(Paragraph(f"Document: {report_data['filename']}", styles['Normal']))
            elements.append(Paragraph(f"Generated: {report_data['generated_at']}", styles['Normal']))
            elements.append(Paragraph(f"Confidence: {report_data['confidence_level']:.1%}", styles['Normal']))
            elements.append(Spacer(1, 0.3 * inch))
            
            # Objective
            elements.append(Paragraph("Objective", heading_style))
            elements.append(Paragraph(report_data['objective'], styles['Normal']))
            elements.append(Spacer(1, 0.2 * inch))
            
            # Challenges
            elements.append(Paragraph("Current Challenges", heading_style))
            for i, challenge in enumerate(report_data['challenges'], 1):
                elements.append(Paragraph(f"{i}. {challenge}", styles['Normal']))
            elements.append(Spacer(1, 0.2 * inch))
            
            # Proposed Solutions
            elements.append(Paragraph("Proposed Solutions", heading_style))
            for i, solution in enumerate(report_data['proposed_solutions'], 1):
                elements.append(Paragraph(f"{i}. {solution}", styles['Normal']))
            elements.append(Spacer(1, 0.2 * inch))
            
            # Key Insights
            elements.append(Paragraph("Key Insights", heading_style))
            for i, insight in enumerate(report_data['key_insights'], 1):
                elements.append(Paragraph(f"• {insight}", styles['Normal']))
            elements.append(Spacer(1, 0.3 * inch))
            
            # Conversation Summary
            elements.append(PageBreak())
            elements.append(Paragraph("Conversation Summary", heading_style))
            
            for exchange in report_data['conversation_history'][-5:]:  # Last 5 exchanges
                if exchange['role'] == 'user':
                    elements.append(Paragraph(f"<b>User:</b> {exchange['content']}", styles['Normal']))
                else:
                    elements.append(Paragraph(f"<b>Assistant:</b> {exchange['content'][:200]}...", styles['Normal']))
                elements.append(Spacer(1, 0.1 * inch))
            
            # Build PDF
            doc.build(elements)
            return str(filename)
        
        except Exception as e:
            print(f"Error generating PDF: {e}")
            raise
    
    def _generate_docx_report(self, report_data: Dict[str, Any]) -> str:
        """Generate DOCX report."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = self.output_dir / f"Report_{timestamp}.docx"
            
            # Create document
            doc = Document()
            
            # Title
            title = doc.add_heading('Analysis Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Document Info
            info_para = doc.add_paragraph()
            info_para.add_run(f"Document: ").bold = True
            info_para.add_run(report_data['filename'])
            
            info_para = doc.add_paragraph()
            info_para.add_run(f"Generated: ").bold = True
            info_para.add_run(report_data['generated_at'])
            
            info_para = doc.add_paragraph()
            info_para.add_run(f"Confidence: ").bold = True
            info_para.add_run(f"{report_data['confidence_level']:.1%}")
            
            doc.add_paragraph()
            
            # Objective
            doc.add_heading('Objective', 1)
            doc.add_paragraph(report_data['objective'])
            
            # Challenges
            doc.add_heading('Current Challenges', 1)
            for challenge in report_data['challenges']:
                doc.add_paragraph(challenge, style='List Bullet')
            
            # Proposed Solutions
            doc.add_heading('Proposed Solutions', 1)
            for solution in report_data['proposed_solutions']:
                doc.add_paragraph(solution, style='List Bullet')
            
            # Key Insights
            doc.add_heading('Key Insights', 1)
            for insight in report_data['key_insights']:
                doc.add_paragraph(insight, style='List Bullet')
            
            # Conversation Summary
            doc.add_page_break()
            doc.add_heading('Conversation Summary', 1)
            
            for exchange in report_data['conversation_history'][-5:]:
                para = doc.add_paragraph()
                if exchange['role'] == 'user':
                    para.add_run('User: ').bold = True
                else:
                    para.add_run('Assistant: ').bold = True
                para.add_run(exchange['content'][:200] + "..." if len(exchange['content']) > 200 else exchange['content'])
            
            # Save
            doc.save(str(filename))
            return str(filename)
        
        except Exception as e:
            print(f"Error generating DOCX: {e}")
            raise
    
    def generate_json_report(self, report_data: Dict[str, Any]) -> str:
        """Generate JSON report."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = self.output_dir / f"Report_{timestamp}.json"
        
        with open(filename, 'w') as f:
            json.dump(report_data, f, indent=2)
        
        return str(filename)
